import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

ROOT_DIR = r"C:\Users\andre\projects\railway-pg-test"


EXCLUDE_DIRS = {'venv', '__pycache__', '.git', '.mypy_cache', '.pytest_cache'}
EXCLUDE_SUFFIXES = {'.dist-info', '.egg-info'}
EXCLUDE_FILES = {'.env'}
INCLUDE_EXTENSIONS = {'.py', '.md', '.txt', '.html', '.css', '.js'}  # Excludes JSON

def should_exclude_dir(name):
    return name in EXCLUDE_DIRS or any(name.endswith(suffix) for suffix in EXCLUDE_SUFFIXES)

def should_include_file(entry):
    ext = os.path.splitext(entry)[1]
    return ext in INCLUDE_EXTENSIONS and entry not in EXCLUDE_FILES

def icon_for(entry, is_dir):
    if is_dir:
        return "📁"
    elif entry.endswith(".py"):
        return "🐍"
    else:
        return "📄"

def read_file_content(filepath):
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.readlines()
    except:
        return None

def add_code_block(doc, lines):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    for i, line in enumerate(lines):
        run = para.add_run(line if line.endswith('\n') else line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)  # #333333
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

def apply_folder_style(paragraph):
    if not paragraph.runs:
        return
    run = paragraph.runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

def write_project_map(root_dir, output_docx):
    doc = Document()
    doc.add_heading("📁 Folder & File Structure", level=1)

    folder_map_lines = []
    code_blocks = []

    def scan_structure(path, level=0):
        try:
            entries = sorted(os.listdir(path))
        except PermissionError:
            return
        for entry in entries:
            if should_exclude_dir(entry):
                continue
            full_path = os.path.join(path, entry)
            is_dir = os.path.isdir(full_path)
            indent = "    " * level
            icon = icon_for(entry, is_dir)
            line = f"{indent}{icon} {entry}/" if is_dir else f"{indent}{icon} {entry}"
            folder_map_lines.append(line)

            if not is_dir and should_include_file(entry):
                content = read_file_content(full_path)
                if content:
                    rel_path = os.path.relpath(full_path, root_dir)
                    code_blocks.append((rel_path, content))

            if is_dir:
                scan_structure(full_path, level + 1)

    folder_map_lines.append(f"{os.path.basename(ROOT_DIR)}/")
    scan_structure(ROOT_DIR, level=1)

    # Folder tree
    for line in folder_map_lines:
        p = doc.add_paragraph(line)
        apply_folder_style(p)

    doc.add_paragraph()  # spacer

    # Code blocks
    doc.add_heading("🧠 Code & Content", level=1)
    for rel_path, lines in code_blocks:
        doc.add_paragraph(f"📄 {rel_path}", style='Heading3')
        add_code_block(doc, lines)

    doc.save(output_docx)
    print(f"✅ Saved clean layout to: {output_docx}")

if __name__ == "__main__":
    write_project_map(ROOT_DIR, OUTPUT_DOCX)
