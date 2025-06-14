from docx import Document
from docx.shared import Pt

OUTPUT_DOCX = "SmartCoach_GPT_Handoff_Summary_CLEAN.docx"

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    doc.add_paragraph(text)

def generate_gpt_handoff_summary():
    doc = Document()

    add_heading(doc, "✅ SmartCoach – GPT Handoff Summary", level=1)
    add_paragraph(doc, "This document was generated to bootstrap a Custom GPT session with complete project context.")

    # Last Known Good State
    add_heading(doc, "🏁 Last known good state:", level=2)
    add_paragraph(doc, "• Tag: v1.0.0")
    add_paragraph(doc, "• Deployment: Railway (web-production-c4329.up.railway.app)")
    add_paragraph(doc, "• OAuth tested: ✅")
    add_paragraph(doc, "• DB Connected: ✅")

    # Environment
    add_heading(doc, "🔐 Environment Settings:", level=2)
    add_paragraph(doc, "• ADMIN_USER=admin")
    add_paragraph(doc, "• REDIRECT_URI=https://web-production-c4329.up.railway.app/oauth/callback")
    add_paragraph(doc, "• DATABASE_URL=(set in Railway)")
    add_paragraph(doc, "• SECRET_KEY, INTERNAL_API_KEY, CRON_SECRET_KEY defined")

    # Architecture
    add_heading(doc, "🔧 Architecture Notes:", level=2)
    add_paragraph(doc, "• Flask API deployed on Railway")
    add_paragraph(doc, "• PostgreSQL managed via Railway plugin")
    add_paragraph(doc, "• JWT-based authentication and user sessions")
    add_paragraph(doc, "• Strava OAuth 2.0 integration and webhook registration planned")

    # Coaching Intelligence Objective
    add_heading(doc, "🎯 Coaching Intelligence Objectives:", level=2)
    add_paragraph(doc, "• Build personalized weekly training plans")
    add_paragraph(doc, "• Track Strava-recorded performance")
    add_paragraph(doc, "• Monitor deviations and adapt based on user behavior")
    add_paragraph(doc, "• Support natural conversation for advising and adjusting")

    # System Components
    add_heading(doc, "🧱 System Components:", level=2)
    add_paragraph(doc, "• Custom GPT logic: planning, feedback, and conversation")
    add_paragraph(doc, "• PostgreSQL DB storing plan and activity data")
    add_paragraph(doc, "• Strava Sync service")
    add_paragraph(doc, "• Planned: Training plan generator, weekly comparator")

    # Known Issues
    add_heading(doc, "⚠️ Known Issues:", level=2)
    add_paragraph(doc, "• None at last deploy. Strava activity fetch untested.")

    # Tasks
    add_heading(doc, "📌 Next Tasks:", level=2)
    add_paragraph(doc, "1. Test full Strava flow: code -> token -> activity fetch")
    add_paragraph(doc, "2. Implement auto-sync")
    add_paragraph(doc, "3. Add athlete-to-user ID mapping")
    add_paragraph(doc, "4. Connect activity analysis")
    add_paragraph(doc, "5. Ship Phase 2 readiness")

    # GPT Boot Prompt
    add_heading(doc, "🧠 GPT Boot Prompt", level=2)
    add_paragraph(doc, "Create a Custom GPT with the following context:")
    add_paragraph(doc, "- Name: SmartCoachDev")
    add_paragraph(doc, "- Role: Flask engineer + PM for Smart Marathon Coach API")
    add_paragraph(doc, "- Accesses: JWT auth, Strava OAuth, Railway deploy logs, pg DB")
    add_paragraph(doc, "- Code base starts in: /src, /templates, /scripts")
    add_paragraph(doc, "- Current live env: https://web-production-c4329.up.railway.app")
    add_paragraph(doc, "- Start by reviewing recent work, then continue with task #1: Test full Strava activity flow.")
    add_paragraph(doc, "- Reference: ‘Smart Coach - Project Reference.docx’ for architecture, history, and schema.")


    from docx import Document as DocxDocument

    def append_docx_content(target_doc, source_path):
        source_doc = DocxDocument(source_path)
        for element in source_doc.element.body:
            target_doc.element.body.append(element)

    # At the end of generate_gpt_handoff_summary():
        try:
            append_docx_content(doc, "final_project_map.docx")
            print("✅ Appended folder/code map from final_project_map.docx")
        except Exception as e:
            print(f"⚠️ Could not append code map: {e}")






    doc.save(OUTPUT_DOCX)
    print(f"✅ Summary written to {OUTPUT_DOCX}")

if __name__ == "__main__":
    generate_gpt_handoff_summary()
